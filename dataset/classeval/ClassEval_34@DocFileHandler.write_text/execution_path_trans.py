
import sys
import trace
import os
import unittest
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path

    def read_text(self):
        doc = Document(self.file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)

    def write_text(self, content, font_size=12, alignment='left'):
        try:
            doc = Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(content)
            font = run.font
            font.size = Pt(font_size)
            alignment_value = self._get_alignment_value(alignment)
            paragraph.alignment = alignment_value
            doc.save(self.file_path)
            return True
        except:
            return False

    def add_heading(self, heading, level=1):
        try:
            doc = Document(self.file_path)
            doc.add_heading(heading, level)
            doc.save(self.file_path)
            return True
        except:
            return False

    def add_table(self, data):
        try:
            doc = Document(self.file_path)
            table = doc.add_table(rows=len(data), cols=len(data[0]))
            for i, row in enumerate(data):
                for j, cell_value in enumerate(row):
                    table.cell(i, j).text = str(cell_value)
            doc.save(self.file_path)
            return True
        except:
            return False

    def _get_alignment_value(self, alignment):
        alignment_options = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_options.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
class Test(unittest.TestCase):
    def test(self):
            self.file_path = "test_example.docx"
            self.handler = DocFileHandler(self.file_path)
            doc = Document()
            doc.add_paragraph("Initial content")
            doc.save(self.file_path)
            new_content = "New content 1"
            self.handler.write_text(new_content)
            text_content = self.handler.read_text()
            return text_content,new_content
t=Test()
a=t.test()
with open('/home/changshu/CODEMIND/dataset/classeval/ClassEval_34@DocFileHandler.write_text/output.txt', 'w') as wr:
    wr.write(str(a))
        

tracer = trace.Trace(
    ignoredirs=[sys.prefix, sys.exec_prefix],
    trace=1,
    count=0)
tracer.run('t.test()')
